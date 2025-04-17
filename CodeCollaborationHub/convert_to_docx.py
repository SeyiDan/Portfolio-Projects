import subprocess
import os

def markdown_to_docx(md_file, docx_file):
    """Convert a markdown file to docx using pandoc"""
    try:
        # Try to use pandoc if available
        subprocess.run(['pandoc', md_file, '-o', docx_file], check=True)
        print(f"Successfully converted {md_file} to {docx_file}")
        return True
    except (subprocess.SubprocessError, FileNotFoundError):
        print("Pandoc conversion failed or not available")
        try:
            # Alternative method using python-docx
            import markdown
            from docx import Document
            from docx.shared import Pt
            
            # Read markdown file
            with open(md_file, 'r', encoding='utf-8') as f:
                md_content = f.read()
            
            # Convert markdown to HTML
            html = markdown.markdown(md_content)
            
            # Create a new Word document
            doc = Document()
            
            # Add title
            title = "CodeCollaborationHub Documentation"
            doc.add_heading(title, 0)
            
            # Add paragraphs
            lines = md_content.split('\n')
            for line in lines:
                if line.startswith('# '):
                    # Level 1 heading
                    doc.add_heading(line[2:], 1)
                elif line.startswith('## '):
                    # Level 2 heading
                    doc.add_heading(line[3:], 2)
                elif line.startswith('### '):
                    # Level 3 heading
                    doc.add_heading(line[4:], 3)
                elif line.startswith('#### '):
                    # Level 4 heading
                    doc.add_heading(line[5:], 4)
                elif line.startswith('- ') or line.startswith('* '):
                    # List item
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line.startswith('1. ') or line.startswith('2. '):
                    # Numbered list item
                    doc.add_paragraph(line[3:], style='List Number')
                elif not line.strip():
                    # Empty line
                    doc.add_paragraph()
                elif line.startswith('```'):
                    # Code block - just add as plain text for simplicity
                    continue
                else:
                    # Regular paragraph
                    doc.add_paragraph(line)
            
            # Save the document
            doc.save(docx_file)
            print(f"Successfully created {docx_file} using python-docx")
            return True
        except Exception as e:
            print(f"Failed to create Word document: {e}")
            return False

if __name__ == "__main__":
    md_file = "Documentation.md"
    docx_file = "CodeCollaborationHub_Documentation.docx"
    
    if os.path.exists(md_file):
        markdown_to_docx(md_file, docx_file)
    else:
        print(f"Error: {md_file} not found") 